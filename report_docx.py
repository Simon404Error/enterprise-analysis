# -*- coding: utf-8 -*-
"""
Word 报告生成模块（可开关 · 自由排序 · 样式可选）
=============================================

- 报告由若干「模块」组成，每个模块可：启用/关闭、调整顺序（自由排序）、
  自选呈现样式（表格 / 柱状图 / 折线图）。
- 默认按用户要求：
    1) 不显示 纳税总额、负债总额、资产负债率、销售净利率、净利润率、环比增长率；
       企业公示分布默认不分析；增长率用折线图；
    2) 着重文字分析（每个模块先给文字分析段落，不是满屏图表）；
    3) 涉及净资产收益率(ROE)、总资产周转率时附文字解释，说明其反映经营哪些方面；
    4) 对存在增长趋势的地区，着重分析其增长率、增长趋势与（定性）增长原因；
    5) 报告不含「建议」章节。
- 两种方式：
    方式A：直接依据 data/ 计算结果自动生成。
    方式B：先生成「报告填写表单.xlsx」（含报告设置/字段显示/封面/数值），填写后生成。
运行：python make_report.py -a  /  -f  /  表单.xlsx
"""
import os
import json
import shutil
import tempfile

import numpy as np
import pandas as pd

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'SimSun']
plt.rcParams['axes.unicode_minus'] = False

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from config import (COL_REGION, COL_YEAR,
                    REPORT_DEFAULT_DOCX, REPORT_FORM_XLSX,
                    REPORT_IMG_DIR, REPORT_IMG_MAP, REPORT_SETTINGS_JSON,
                    PUBLISH_ALL, PUBLISH_PART, PUBLISH_NONE, GROWTH_COLS)
from analysis import (region_yearly, add_region_metrics, publish_by_region_year,
                      metrics_for_display)

# ---------------- 模块与默认配置 ----------------
STYLES = ['表格', '柱状图', '折线图']

DEFAULT_MODULES = [
    {'id': 'overall', 'name': '生产经营总体情况', 'enabled': True, 'order': 1, 'style': '表格'},
    {'id': 'metrics', 'name': '经营指标分析',     'enabled': True, 'order': 2, 'style': '表格'},
    {'id': 'growth',  'name': '增长率与增长趋势分析', 'enabled': True, 'order': 3, 'style': '折线图'},
    {'id': 'publish', 'name': '企业公示分布',     'enabled': False, 'order': 4, 'style': '柱状图'},
]

# 中文序号（按模块启用后的顺序动态编号，避免乱序）
_CN_NUM = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']


def _numbered_title(idx, name):
    """按顺序生成标题：一、XXX / 二、XXX ..."""
    num = _CN_NUM[idx] if idx < len(_CN_NUM) else f'{idx + 1}.'
    return f'{num}、{name}'

# 默认不显示的列（可经「报告设置.json」/ 表单「字段显示」/ 网页勾选调整）
DEFAULT_HIDE = ['负债总额', '纳税总额', '资产负债率', '销售净利率', '净利润率']
# 增长率：默认只显示同比
ALL_GROWTH_VISIBLE = ['销售额同比增长率', '净利润同比增长率',
                      '资产总额同比增长率', '主营业务收入同比增长率']

# 图表名称（自定义图片映射用）
CHART_M1 = '净利率ROE趋势'
CHART_M2 = '各地区盈利对比'
CHART_M3 = '销售额净利润趋势'
CHART_P1 = '公示分布按年'
CHART_NAMES = [CHART_M1, CHART_M2, CHART_M3, CHART_P1]


# ============================ 配置 ============================
def load_config():
    """读取模块/字段配置：优先「报告设置.json」，否则默认。返回 (modules, hide_cols)。"""
    modules = [dict(m) for m in DEFAULT_MODULES]
    hide = list(DEFAULT_HIDE)
    if os.path.isfile(REPORT_SETTINGS_JSON):
        try:
            with open(REPORT_SETTINGS_JSON, encoding='utf-8') as f:
                cfg = json.load(f)
            if isinstance(cfg.get('modules'), list):
                by_id = {m['id']: m for m in modules}
                for m in cfg['modules']:
                    mid = m.get('id')
                    if mid in by_id:
                        by_id[mid].update(
                            {k: m[k] for k in ('enabled', 'order', 'style') if k in m})
            if isinstance(cfg.get('hide'), list):
                hide = list(cfg['hide'])
        except Exception:
            pass
    return modules, hide


def _visible_growth(hide):
    return [c for c in ALL_GROWTH_VISIBLE if c not in hide]


# ============================ 数据 ============================
def _load_analysis():
    import file_loader
    detail, msgs = file_loader.load_all()
    if detail.empty:
        return detail, pd.DataFrame(), pd.DataFrame()
    region = add_region_metrics(region_yearly(detail))
    _, pub = publish_by_region_year(detail)
    return detail, region, pub


def _fmt(v, digits=2):
    if v is None or pd.isna(v):
        return 'NA'
    return f'{v:,.{digits}f}'


def _overall_table(detail, region, hide):
    cnt = detail.groupby([COL_REGION, COL_YEAR]).size().rename('企业数量').reset_index()
    rows = []
    for _, r in region.sort_values([COL_REGION, COL_YEAR]).iterrows():
        c = cnt[(cnt[COL_REGION] == r[COL_REGION]) & (cnt[COL_YEAR] == r[COL_YEAR])]
        row = {COL_REGION: r[COL_REGION], COL_YEAR: str(r[COL_YEAR]),
               '企业数量': int(c['企业数量'].iloc[0]) if not c.empty else 0}
        for k in ['资产总额', '销售额或营业收入', '净利润']:
            if k not in hide:
                row[k] = r[k]
        rows.append(row)
    return pd.DataFrame(rows)


def _metrics_table(region, hide):
    m = metrics_for_display(region)
    cols = [c for c in ['净资产收益率(ROE)', '主营业务收入占比', '总资产周转率',
                        '销售净利率', '净利润率', '资产负债率'] if c not in hide]
    keep = [COL_REGION, COL_YEAR] + cols
    return m[[c for c in keep if c in m.columns]].copy()


def _growth_table(region, hide, growth_cols):
    m = metrics_for_display(region)
    keep = [c for c in growth_cols if c in m.columns]
    return m[[COL_REGION, COL_YEAR] + keep].copy() if keep else None


# ============================ 叙述文字 ============================
def _overall_narrative(detail, region):
    if region.empty:
        return ''
    years = sorted(region[COL_YEAR].unique())
    nf = (f'本报告基于「年度报告记录」数据，共覆盖 {region[COL_REGION].nunique()} 个地区、'
          f'{len(years)} 个年度（{years[0]}~{years[-1]}），合计 {len(detail)} 条企业年度记录。')
    tot = region[['资产总额', '销售额或营业收入', '净利润']].sum()
    parts = [f'全区全年合计：资产总额 {_fmt(tot["资产总额"])} 万元，销售额（营业收入）'
             f'{_fmt(tot["销售额或营业收入"])} 万元，净利润 {_fmt(tot["净利润"])} 万元。']
    if tot['销售额或营业收入'] != 0:
        parts.append(f'全区整体销售净利率为 {_fmt(tot["净利润"] / tot["销售额或营业收入"] * 100)}%。')
    return nf + ''.join(parts)


_ROE_EXPLAIN = ('净资产收益率(ROE) = 净利润 / 所有者权益合计，衡量企业运用所有者（股东）投入资本'
                '创造净利润的能力，反映企业为股东带来回报的水平；ROE 越高，说明资本利用效率越高、'
                '股东回报越好，是企业盈利能力与资本配置效率的关键体现。')
_TURNOVER_EXPLAIN = ('总资产周转率 = 营业收入 / 平均资产总额，反映企业全部资产从投入到产出的运营效率，'
                     '即每一元资产能创造多少营业收入；周转率越高，说明资产运营活跃、产能利用率越高，'
                     '反之说明资产闲置或经营放缓，是衡量企业资产运营效率的重要指标。')


def _metrics_narrative(region):
    if region.empty:
        return ''
    latest = region[COL_YEAR].max()
    rows = region[region[COL_YEAR] == latest]
    m = metrics_for_display(rows)
    lines = [f'以下为 {latest} 年度各地区经营指标（比率为 %，NA 表示分母为 0/缺失无法计算）。']
    if '净资产收益率(ROE)' in m and m['净资产收益率(ROE)'].notna().any():
        bi = m['净资产收益率(ROE)'].idxmax()
        lines.append(f'ROE 最高的地区为 {m.loc[bi, COL_REGION]}'
                     f'（{_fmt(m.loc[bi, "净资产收益率(ROE)"])}%）。' + _ROE_EXPLAIN)
    if '总资产周转率' in m and m['总资产周转率'].notna().any():
        ti = m['总资产周转率'].idxmax()
        lines.append(f'总资产周转率最高的地区为 {m.loc[ti, COL_REGION]}'
                     f'（{_fmt(m.loc[ti, "总资产周转率"])}%）。' + _TURNOVER_EXPLAIN)
    if '主营业务收入占比' in m and m['主营业务收入占比'].notna().any():
        lines.append('主营业务收入占比反映企业对主营业务的依赖程度，占比高说明经营围绕主业、结构相对稳定。')
    return ''.join(lines)


def _growth_narrative(region, growth_cols):
    """对存在增长趋势的地区，着重分析其增长率、增长趋势与（定性）增长原因。"""
    if region.empty:
        return ''
    labels = {'销售额同比增长率': '销售额', '净利润同比增长率': '净利润',
              '资产总额同比增长率': '资产总额',
              '主营业务收入同比增长率': '主营业务收入'}
    parts = ['以下按地区分析各指标同比（严格相邻年度计算，上年缺失/为 0 记 NA）。']
    for reg, g in region.groupby(COL_REGION):
        g = g.sort_values(COL_YEAR)
        items = []
        for c in growth_cols:
            if c not in g.columns:
                continue
            v = g[c].iloc[-1]
            if pd.notna(v):
                items.append((labels.get(c, c), float(v)))
        if not items:
            continue
        desc = '、'.join(f'{k}同比 {_fmt(v * 100)}%' for k, v in items)
        trend = ('总体呈增长趋势' if any(v > 0 for _, v in items)
                 else ('总体呈下降趋势' if any(v < 0 for _, v in items) else '总体基本持平'))
        causes = []
        kv = dict(items)
        if kv.get('销售额', 0) > 0:
            causes.append('销售额增长通常反映需求扩大与市场开拓')
        if kv.get('资产总额', 0) > 0:
            causes.append('资产总额增长反映资本投入扩大、规模扩张')
        if kv.get('净利润', 0) > 0:
            causes.append('净利润增长反映盈利改善与利润增厚')
        if kv.get('销售额', 0) < 0:
            causes.append('销售额下降可能反映需求放缓或竞争加剧')
        if kv.get('净利润', 0) < 0:
            causes.append('净利润下降提示盈利承压')
        reason = ('，可能原因：' + '、'.join(causes) + '。') if causes else '。'
        parts.append(f'【{reg}】{desc}，{trend}{reason}')
    return ''.join(parts)


def _publish_narrative(pub):
    if pub.empty:
        return ''
    lines = ['按企业年度记录的 12 项判定公示状况（数值为 0 视为空位）。']
    for _, r in pub.sort_values([COL_REGION, COL_YEAR]).iterrows():
        e = (int(r.get(PUBLISH_ALL, 0)) + int(r.get(PUBLISH_PART, 0))
             + int(r.get(PUBLISH_NONE, 0)))
        if e == 0:
            continue
        lines.append(f'{r[COL_REGION]} {r[COL_YEAR]} 年共 {e} 家企业：全部公示 '
                     f'{int(r.get(PUBLISH_ALL, 0))} 家、部分公示 {int(r.get(PUBLISH_PART, 0))} 家、'
                     f'全部不公示 {int(r.get(PUBLISH_NONE, 0))} 家。')
    return ''.join(lines)


# ============================ 图表生成 ============================
def _chart_overall(region, style, img_dir):
    if region.empty:
        return None
    if style == '柱状图':
        d = region.copy()
        d['label'] = d[COL_REGION].astype(str) + '·' + d[COL_YEAR].astype(str)
        fig, ax = plt.subplots(figsize=(7.2, 3.4))
        x = np.arange(len(d)); w = 0.36
        ax.bar(x - w / 2, d['销售额或营业收入'], w, label='销售额')
        ax.bar(x + w / 2, d['净利润'], w, label='净利润')
        ax.set_xticks(x); ax.set_xticklabels(d['label'], rotation=40, ha='right', fontsize=8)
        ax.set_ylabel('万元'); ax.set_title('各年度各地区 销售额 / 净利润')
        ax.legend(); ax.grid(True, axis='y', alpha=0.3)
        fig.tight_layout()
        p = os.path.join(img_dir, '总体_柱状图.png')
        fig.savefig(p, dpi=150); plt.close(fig)
        return p
    if style == '折线图':
        g = region.groupby(COL_YEAR).agg(销售额=('销售额或营业收入', 'sum'),
                                         净利润=('净利润', 'sum')).reset_index()
        fig, ax = plt.subplots(figsize=(7.2, 3.4))
        ax.plot(g[COL_YEAR], g['销售额'], marker='o', label='销售额')
        ax.plot(g[COL_YEAR], g['净利润'], marker='s', label='净利润')
        ax.set_xlabel('年度'); ax.set_ylabel('万元'); ax.set_title('销售额与净利润趋势（万元）')
        ax.legend(); ax.grid(True, alpha=0.3)
        fig.tight_layout()
        p = os.path.join(img_dir, '总体_折线图.png')
        fig.savefig(p, dpi=150); plt.close(fig)
        return p
    return None


def _chart_metrics(region, style, img_dir):
    if region.empty:
        return None
    if style == '柱状图':
        latest = region[COL_YEAR].max()
        bar = metrics_for_display(region[region[COL_YEAR] == latest]).sort_values(COL_REGION)
        x = bar[COL_REGION].astype(str)
        fig, ax = plt.subplots(figsize=(7.2, 3.4))
        idx = np.arange(len(x)); w = 0.27
        for j, c in enumerate(['净资产收益率(ROE)', '主营业务收入占比', '总资产周转率']):
            if c in bar.columns:
                ax.bar(idx + (j - 1) * w, bar[c], w, label=c)
        ax.set_xticks(idx); ax.set_xticklabels(x, rotation=30, ha='right')
        ax.set_ylabel('%'); ax.set_title(f'{latest} 年各地区经营指标对比')
        ax.legend(); ax.grid(True, axis='y', alpha=0.3)
        fig.tight_layout()
        p = os.path.join(img_dir, '指标_柱状图.png')
        fig.savefig(p, dpi=150); plt.close(fig)
        return p
    if style == '折线图':
        g = region.groupby(COL_YEAR).agg(净利=('净利润', 'sum'), 权益=('所有者权益合计', 'sum'),
                                         销售=('销售额或营业收入', 'sum'),
                                         资产=('资产总额', 'sum')).reset_index()
        g['ROE'] = g['净利'] / g['权益'].replace(0, np.nan) * 100
        g['周转'] = g['销售'] / g['资产'].replace(0, np.nan) * 100
        fig, ax = plt.subplots(figsize=(7.2, 3.4))
        ax.plot(g[COL_YEAR], g['ROE'], marker='o', label='净资产收益率(ROE)%')
        ax.plot(g[COL_YEAR], g['周转'], marker='s', label='总资产周转率%')
        ax.set_xlabel('年度'); ax.set_ylabel('%')
        ax.set_title('地区整体 ROE / 总资产周转率 趋势')
        ax.legend(); ax.grid(True, alpha=0.3)
        fig.tight_layout()
        p = os.path.join(img_dir, '指标_折线图.png')
        fig.savefig(p, dpi=150); plt.close(fig)
        return p
    return None


def _chart_growth(region, growth_cols, style, img_dir):
    if region.empty or not growth_cols:
        return None
    if style == '折线图':
        n = len(growth_cols)
        fig, axes = plt.subplots(2, 2, figsize=(9.2, 6.4)) if n > 1 else \
            plt.subplots(1, 1, figsize=(9.2, 3.4))
        axes = np.array(axes).ravel()
        for ax, c in zip(axes, growth_cols):
            for reg, g in region.groupby(COL_REGION):
                g = g.sort_values(COL_YEAR)
                vals = metrics_for_display(g)
                if c in vals.columns:
                    ax.plot(g[COL_YEAR], vals[c], marker='o', label=str(reg))
            ax.set_title(c.replace('同比增长率', '同比(%)'))
            ax.legend(fontsize=7); ax.grid(True, alpha=0.3)
        for ax in axes[len(growth_cols):]:
            ax.set_visible(False)
        fig.suptitle('各指标同比增长率（折线）')
        fig.tight_layout(rect=[0, 0, 1, 0.97])
        p = os.path.join(img_dir, '增长_折线图.png')
        fig.savefig(p, dpi=150); plt.close(fig)
        return p
    if style == '柱状图':
        latest = region[COL_YEAR].max()
        bar = metrics_for_display(region[region[COL_YEAR] == latest]).sort_values(COL_REGION)
        x = bar[COL_REGION].astype(str)
        fig, ax = plt.subplots(figsize=(7.2, 3.4))
        idx = np.arange(len(x)); w = 0.2
        for j, c in enumerate(growth_cols):
            if c in bar.columns:
                ax.bar(idx + (j - len(growth_cols) / 2 + 0.5) * w, bar[c], w,
                       label=c.replace('同比增长率', ''))
        ax.axhline(0, color='grey', lw=0.8)
        ax.set_xticks(idx); ax.set_xticklabels(x, rotation=30, ha='right')
        ax.set_ylabel('%'); ax.set_title(f'{latest} 年各地区同比增长率对比')
        ax.legend(fontsize=7); ax.grid(True, axis='y', alpha=0.3)
        fig.tight_layout()
        p = os.path.join(img_dir, '增长_柱状图.png')
        fig.savefig(p, dpi=150); plt.close(fig)
        return p
    return None


def _form_chart(df, mid, style, img_dir, growth_cols=None):
    """按表单数据出图（样式可选）：overall/metrics/growth/publish。返回图片路径或 None。"""
    if df is None or df.empty or style not in ('柱状图', '折线图'):
        return None
    try:
        if mid == 'overall':
            if style == '柱状图':
                d = df.copy()
                d['label'] = d[COL_REGION].astype(str) + '·' + d[COL_YEAR].astype(str)
                fig, ax = plt.subplots(figsize=(7.2, 3.4))
                x = np.arange(len(d)); w = 0.36
                ax.bar(x - w / 2, d['销售额或营业收入'], w, label='销售额')
                ax.bar(x + w / 2, d['净利润'], w, label='净利润')
                ax.set_xticks(x); ax.set_xticklabels(d['label'], rotation=40, ha='right', fontsize=8)
                ax.set_ylabel('万元'); ax.set_title('各年度各地区 销售额 / 净利润')
                ax.legend(); ax.grid(True, axis='y', alpha=0.3)
            else:
                g = df.groupby(COL_YEAR).agg(销售额=('销售额或营业收入', 'sum'),
                                             净利润=('净利润', 'sum')).reset_index()
                fig, ax = plt.subplots(figsize=(7.2, 3.4))
                ax.plot(g[COL_YEAR], g['销售额'], marker='o', label='销售额')
                ax.plot(g[COL_YEAR], g['净利润'], marker='s', label='净利润')
                ax.set_xlabel('年度'); ax.set_ylabel('万元')
                ax.set_title('销售额与净利润趋势（万元）')
                ax.legend(); ax.grid(True, alpha=0.3)
        elif mid == 'metrics':
            if style == '柱状图':
                latest = df[COL_YEAR].max()
                bar = df[df[COL_YEAR] == latest].sort_values(COL_REGION)
                x = bar[COL_REGION].astype(str)
                fig, ax = plt.subplots(figsize=(7.2, 3.4))
                idx = np.arange(len(x)); w = 0.27
                for j, c in enumerate(['净资产收益率(ROE)', '主营业务收入占比', '总资产周转率']):
                    if c in bar.columns:
                        ax.bar(idx + (j - 1) * w, bar[c], w, label=c)
                ax.set_xticks(idx); ax.set_xticklabels(x, rotation=30, ha='right')
                ax.set_ylabel('%'); ax.set_title(f'{latest} 年各地区经营指标对比')
                ax.legend(); ax.grid(True, axis='y', alpha=0.3)
            else:
                g = df.groupby(COL_YEAR).mean(numeric_only=True).reset_index()
                fig, ax = plt.subplots(figsize=(7.2, 3.4))
                if '净资产收益率(ROE)' in g:
                    ax.plot(g[COL_YEAR], g['净资产收益率(ROE)'], marker='o', label='净资产收益率(ROE)%')
                if '总资产周转率' in g:
                    ax.plot(g[COL_YEAR], g['总资产周转率'], marker='s', label='总资产周转率%')
                ax.set_xlabel('年度'); ax.set_ylabel('%')
                ax.set_title('经营指标年度趋势（%）')
                ax.legend(); ax.grid(True, alpha=0.3)
        elif mid == 'growth':
            gcols = growth_cols or [c for c in ALL_GROWTH_VISIBLE if c in df.columns]
            if not gcols:
                return None
            if style == '折线图':
                n = len(gcols)
                if n > 1:
                    fig, axes = plt.subplots(2, 2, figsize=(9.2, 6.4))
                else:
                    fig, axes = plt.subplots(1, 1, figsize=(9.2, 3.4))
                axes = np.array(axes).ravel()
                for ax, c in zip(axes, gcols):
                    for reg, g in df.groupby(COL_REGION):
                        g = g.sort_values(COL_YEAR)
                        if c in g:
                            ax.plot(g[COL_YEAR], g[c], marker='o', label=str(reg))
                    ax.set_title(c.replace('同比增长率', '同比(%)'))
                    ax.legend(fontsize=7); ax.grid(True, alpha=0.3)
                for ax in axes[len(gcols):]:
                    ax.set_visible(False)
                fig.suptitle('各指标同比增长率（折线）')
                fig.tight_layout(rect=[0, 0, 1, 0.97])
            else:
                latest = df[COL_YEAR].max()
                bar = df[df[COL_YEAR] == latest].sort_values(COL_REGION)
                x = bar[COL_REGION].astype(str)
                fig, ax = plt.subplots(figsize=(7.2, 3.4))
                idx = np.arange(len(x)); w = 0.2
                for j, c in enumerate(gcols):
                    if c in bar.columns:
                        ax.bar(idx + (j - len(gcols) / 2 + 0.5) * w, bar[c], w,
                               label=c.replace('同比增长率', ''))
                ax.axhline(0, color='grey', lw=0.8)
                ax.set_xticks(idx); ax.set_xticklabels(x, rotation=30, ha='right')
                ax.set_ylabel('%'); ax.set_title(f'{latest} 年各地区同比增长率对比')
                ax.legend(fontsize=7); ax.grid(True, axis='y', alpha=0.3)
        elif mid == 'publish':
            d = df.copy()
            d['label'] = d[COL_REGION].astype(str) + '·' + d[COL_YEAR].astype(str)
            d = d.sort_values([COL_REGION, COL_YEAR])
            fig, ax = plt.subplots(figsize=(7.2, 3.6))
            idx = np.arange(len(d))
            ax.bar(idx, d[PUBLISH_ALL], color='#4C9F70', label=PUBLISH_ALL)
            ax.bar(idx, d[PUBLISH_PART], bottom=d[PUBLISH_ALL], color='#E8A33D', label=PUBLISH_PART)
            ax.bar(idx, d[PUBLISH_NONE], bottom=d[PUBLISH_ALL] + d[PUBLISH_PART],
                   color='#C0504D', label=PUBLISH_NONE)
            ax.set_xticks(idx); ax.set_xticklabels(d['label'], rotation=45, ha='right', fontsize=8)
            ax.set_ylabel('企业数'); ax.set_title('企业公示分布（按地区·年度）')
            ax.legend(); ax.grid(True, axis='y', alpha=0.3)
        else:
            return None
        fig.tight_layout()
        p = os.path.join(img_dir, f'form_{mid}.png')
        fig.savefig(p, dpi=150); plt.close(fig)
        return p
    except Exception:
        try:
            plt.close('all')
        except Exception:
            pass
        return None


def _chart_publish(pub, style, img_dir):
    if pub.empty:
        return None
    d = pub.copy()
    d['label'] = d[COL_REGION].astype(str) + '·' + d[COL_YEAR].astype(str)
    if len(d) > 15:
        latest = sorted(d[COL_YEAR].unique())[-2:]
        d = d[d[COL_YEAR].isin(latest)]
    d = d.sort_values([COL_REGION, COL_YEAR])
    fig, ax = plt.subplots(figsize=(7.2, 3.6))
    idx = np.arange(len(d))
    ax.bar(idx, d[PUBLISH_ALL], color='#4C9F70', label=PUBLISH_ALL)
    ax.bar(idx, d[PUBLISH_PART], bottom=d[PUBLISH_ALL], color='#E8A33D', label=PUBLISH_PART)
    ax.bar(idx, d[PUBLISH_NONE], bottom=d[PUBLISH_ALL] + d[PUBLISH_PART],
           color='#C0504D', label=PUBLISH_NONE)
    ax.set_xticks(idx); ax.set_xticklabels(d['label'], rotation=45, ha='right', fontsize=8)
    ax.set_ylabel('企业数'); ax.set_title('企业公示分布（按地区·年度）')
    ax.legend(); ax.grid(True, axis='y', alpha=0.3)
    fig.tight_layout()
    p = os.path.join(img_dir, '公示_柱状图.png')
    fig.savefig(p, dpi=150); plt.close(fig)
    return p


# ============================ docx 组装 ============================
def _set_cn(run, name='宋体', size=None, bold=None):
    run.font.name = name
    r = run._element.get_or_add_rPr()
    rF = r.get_or_add_rFonts()
    rF.set(qn('w:eastAsia'), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def _para(doc, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          indent=True, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    p.paragraph_format.space_after = Pt(space_after)
    _set_cn(p.add_run(text), size=size, bold=bold)
    return p


def _heading(doc, text):
    _para(doc, text, size=16, bold=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT,
          space_after=10)


def _add_table(doc, title, headers, rows):
    if title:
        _para(doc, title, size=11, bold=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT,
              space_after=2)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        _set_cn(t.rows[0].cells[j].paragraphs[0].add_run(str(h)), size=10.5, bold=True)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            _set_cn(t.rows[i + 1].cells[j].paragraphs[0].add_run(
                _fmt(v) if isinstance(v, float) else str(v)), size=10.5)
    doc.add_paragraph()


def _add_chart(doc, path):
    if path and os.path.isfile(path):
        doc.add_picture(path, width=Cm(15.5))
        doc.add_paragraph()


def _cover(doc, info):
    for _ in range(6):
        doc.add_paragraph()
    _para(doc, info.get('标题', '年度企业生产经营分析报告'), size=26, bold=True,
          indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    _para(doc, '（' + info.get('地区说明', '全部地区') + ' · ' + info.get('年度说明', '') + '）',
          size=16, bold=True, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    for _ in range(3):
        doc.add_paragraph()
    _para(doc, '编制单位：' + (info.get('编制单位') or info.get('机构', '')), size=14,
          indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    _para(doc, '报告日期：' + (info.get('报告日期') or info.get('日期', '')), size=14,
          indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    doc.add_page_break()


def _default_info(region):
    years = sorted(region[COL_YEAR].unique())
    regions = sorted(region[COL_REGION].unique())
    return {
        '标题': '年度企业生产经营分析报告',
        '地区说明': '全部地区' if len(regions) > 1 else (regions[0] if regions else ''),
        '年度说明': f'{years[0]}~{years[-1]}' if len(years) > 1 else (str(years[0]) if years else ''),
        '报告日期': pd.Timestamp.now().strftime('%Y年%m月%d日'),
    }


_IMG_NAMES = {'overall': '总体情况图', 'metrics': '经营指标图',
              'growth': '增长趋势图', 'publish': '公示分布图'}


def _read_img_override():
    over = {}
    if os.path.isfile(REPORT_IMG_MAP):
        with open(REPORT_IMG_MAP, encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith('#'):
                    continue
                if '=' in line:
                    k, v = line.split('=', 1)
                    over[k.strip()] = v.strip()
    if os.path.isdir(REPORT_IMG_DIR):
        for name in set(_IMG_NAMES.values()):
            if name in over:
                continue
            for ext in ('.png', '.jpg', '.jpeg'):
                p = os.path.join(REPORT_IMG_DIR, name + ext)
                if os.path.isfile(p):
                    over[name] = p
                    break
    return over


def _resolve_override(name, gen_func, overrides):
    p = overrides.get(name)
    if p and os.path.isfile(p):
        return p
    return gen_func() if gen_func else None


def _build(detail, region, pub, info, modules, hide, overrides):
    img_dir = tempfile.mkdtemp(prefix='rep_')
    doc = Document()
    _cover(doc, info)

    growth_cols = _visible_growth(hide)
    data = {
        'overall': _overall_table(detail, region, hide) if not region.empty else pd.DataFrame(),
        'metrics': _metrics_table(region, hide) if not region.empty else pd.DataFrame(),
        'growth': _growth_table(region, hide, growth_cols) if not region.empty else None,
        'publish': pub if not pub.empty else pd.DataFrame(),
    }
    texts = {
        'overall': _overall_narrative(detail, region),
        'metrics': _metrics_narrative(region),
        'growth': _growth_narrative(region, growth_cols),
        'publish': _publish_narrative(pub),
    }
    try:
        for _i, m in enumerate(sorted(modules, key=lambda x: x['order'])):
            mid, style = m['id'], m.get('style', '表格')
            df = data.get(mid)
            txt = texts.get(mid)
            if df is None or df.empty:
                continue
            _heading(doc, _numbered_title(_i, m['name']))
            if txt:
                _para(doc, txt)
            p = None
            if mid == 'overall':
                p = _resolve_override(_IMG_NAMES[mid],
                                      lambda: _chart_overall(region, style, img_dir), overrides)
            elif mid == 'metrics':
                p = _resolve_override(_IMG_NAMES[mid],
                                      lambda: _chart_metrics(region, style, img_dir), overrides)
            elif mid == 'growth':
                p = _resolve_override(_IMG_NAMES[mid],
                                      lambda: _chart_growth(region, growth_cols, style, img_dir),
                                      overrides)
            elif mid == 'publish':
                p = _resolve_override(_IMG_NAMES[mid],
                                      lambda: _chart_publish(pub, style, img_dir), overrides)
            if style == '表格':
                _add_table(doc, '表：' + m['name'], df.columns.tolist(), df.values.tolist())
            else:
                _add_chart(doc, p)
    finally:
        shutil.rmtree(img_dir, ignore_errors=True)

    out = info.get('输出文件') or REPORT_DEFAULT_DOCX
    doc.save(out)
    return out


# ============================ 方式A ============================
def generate_from_data(detail_df=None, region_df=None, info=None, modules=None, hide=None,
                       out_path=REPORT_DEFAULT_DOCX):
    if detail_df is None:
        detail_df, region_df, _ = _load_analysis()
    if region_df is None and not detail_df.empty:
        region_df = add_region_metrics(region_yearly(detail_df))
    _, pub = publish_by_region_year(detail_df) if not detail_df.empty else (None, pd.DataFrame())
    if modules is None or hide is None:
        modules, hide = load_config()
    modules = [dict(m) for m in modules if m.get('enabled', True)]
    modules.sort(key=lambda m: m['order'])
    info = info or _default_info(region_df)
    info['输出文件'] = out_path
    return _build(detail_df, region_df, pub, info, modules, hide, _read_img_override())


# ============================ 方式B：表单 ============================
_FORM_COVER = ['标题', '地区说明', '年度说明', '编制单位', '报告日期']


def generate_form(path=REPORT_FORM_XLSX):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '说明'
    for line in ['【报告填写表单】使用说明', '',
                 '1. 「报告设置」：每模块一行，启用（是/否）、顺序（数字）、呈现样式。',
                 '2. 「字段显示」：字段名一行，填「显示/不显示」；不显示的不进入报告。',
                 '3. 「封面信息」：填写标题、地区说明、年度说明、编制单位、报告日期。',
                 '4. 数值工作表：总体指标、经营指标、增长率、公示分布，按表头填写。',
                 '5. 「图片设置」：填图片路径可自定义对应图表，留空则按样式自动生成。',
                 '6. 填完后运行：python make_report.py 报告填写表单.xlsx',
                 '', '呈现样式：表格 / 柱状图 / 折线图']:
        ws.append([line])

    ws = wb.create_sheet('报告设置')
    ws.append(['模块名称', '启用', '顺序', '呈现样式'])
    for m in DEFAULT_MODULES:
        ws.append([m['name'], '是' if m['enabled'] else '否', m['order'], m['style']])

    ws = wb.create_sheet('字段显示')
    ws.append(['字段名', '是否显示'])
    for c in ['资产总额', '销售额或营业收入', '净利润', '负债总额', '纳税总额',
              '净资产收益率(ROE)', '主营业务收入占比', '总资产周转率',
              '资产负债率', '销售净利率', '净利润率']:
        ws.append([c, '不显示' if c in DEFAULT_HIDE else '显示'])
    for c in ALL_GROWTH_VISIBLE:
        ws.append([c, '显示'])

    ws = wb.create_sheet('封面信息')
    ws.append(['项目', '填写内容'])
    for f in _FORM_COVER:
        ws.append([f, ''])

    ws = wb.create_sheet('图片设置')
    ws.append(['图表名称', '是否自定义', '图片路径'])
    for n in set(_IMG_NAMES.values()):
        ws.append([n, '', ''])

    ws = wb.create_sheet('总体指标')
    ws.append(['地区', '年度', '企业数量', '资产总额', '销售额或营业收入', '净利润'])
    ws = wb.create_sheet('经营指标')
    ws.append(['地区', '年度', '净资产收益率(ROE)', '主营业务收入占比', '总资产周转率'])
    ws = wb.create_sheet('增长率')
    ws.append(['地区', '年度'] + ALL_GROWTH_VISIBLE)
    ws = wb.create_sheet('公示分布')
    ws.append(['地区', '年度', PUBLISH_ALL, PUBLISH_PART, PUBLISH_NONE])

    wb.save(path)
    return path


def _read_form(form_path):
    xl = pd.ExcelFile(form_path)
    modules = [dict(m) for m in DEFAULT_MODULES]
    if '报告设置' in xl.sheet_names:
        df = pd.read_excel(form_path, sheet_name='报告设置')
        name2id = {m['name']: m['id'] for m in DEFAULT_MODULES}
        for _, r in df.iterrows():
            nm = str(r.iloc[0]).strip()
            if nm in name2id:
                en = not ((isinstance(r.iloc[1], str) and str(r.iloc[1]).strip() in ('否', '不', '0'))
                          or (isinstance(r.iloc[1], float) and r.iloc[1] == 0))
                st = str(r.iloc[3]).strip() if pd.notna(r.iloc[3]) else '表格'
                if st not in STYLES:
                    st = '表格'
                for m in modules:
                    if m['id'] == name2id[nm]:
                        m['enabled'] = en
                        if pd.notna(r.iloc[2]):
                            try:
                                m['order'] = int(r.iloc[2])
                            except Exception:
                                pass
                        m['style'] = st
    hide = list(DEFAULT_HIDE)
    if '字段显示' in xl.sheet_names:
        df = pd.read_excel(form_path, sheet_name='字段显示')
        h2 = []
        for _, r in df.iterrows():
            c = str(r.iloc[0]).strip()
            v = str(r.iloc[1]).strip()
            if v in ('不显示', '否', '0'):
                h2.append(c)
        hide = h2
    info = {}
    if '封面信息' in xl.sheet_names:
        df = pd.read_excel(form_path, sheet_name='封面信息')
        for _, r in df.iterrows():
            if pd.notna(r.iloc[0]) and pd.notna(r.iloc[1]):
                info[str(r.iloc[0]).strip()] = str(r.iloc[1]).strip()
    overrides = {}
    if '图片设置' in xl.sheet_names:
        df = pd.read_excel(form_path, sheet_name='图片设置')
        for _, r in df.iterrows():
            if pd.notna(r.iloc[0]) and str(r.iloc[1]).strip().lower() in ('是', 'yes', '1'):
                p = str(r.iloc[2]).strip()
                if p and p.lower() != 'nan' and os.path.isfile(p):
                    overrides[str(r.iloc[0]).strip()] = p

    def _get(sheet):
        if sheet in xl.sheet_names:
            d = pd.read_excel(form_path, sheet_name=sheet)
            return d.dropna(how='all')
        return pd.DataFrame()

    data = {'overall': _get('总体指标'), 'metrics': _get('经营指标'),
            'growth': _get('增长率'), 'publish': _get('公示分布')}
    return info, modules, hide, data, overrides


def _table_intro(df, nm):
    if df is None or df.empty:
        return ''
    return f'以下为「{nm}」数据（金额单位：万元，比率为 %；NA 表示无法计算）。'


def generate_from_form(form_path, out_path=None):
    out_path = out_path or REPORT_DEFAULT_DOCX
    info, modules, hide, data, overrides = _read_form(form_path)
    if not info.get('报告日期'):
        info['报告日期'] = pd.Timestamp.now().strftime('%Y年%m月%d日')
    if not info.get('标题'):
        info['标题'] = '年度企业生产经营分析报告'
    info['输出文件'] = out_path

    img_dir = tempfile.mkdtemp(prefix='rep_')
    doc = Document()
    _cover(doc, info)
    try:
        for _i, m in enumerate(sorted([mm for mm in modules if mm['enabled']],
                                      key=lambda x: x['order'])):
            st = m.get('style', '表格')
            df = data.get(m['id'])
            if df is None or df.empty:
                continue
            _heading(doc, _numbered_title(_i, m['name']))
            _para(doc, _table_intro(df, m['name']))
            if st == '表格':
                _add_table(doc, '表：' + m['name'], df.columns.tolist(), df.values.tolist())
            else:
                p = overrides.get(_IMG_NAMES[m['id']])
                if p and os.path.isfile(p):
                    _add_chart(doc, p)
                else:
                    gcols = _visible_growth(hide) or None
                    p2 = _form_chart(df, m['id'], st, img_dir, gcols)
                    if p2:
                        _add_chart(doc, p2)
                    else:
                        doc.add_paragraph('（该模块数据不足，未插入图表。）')
    finally:
        shutil.rmtree(img_dir, ignore_errors=True)

    doc.save(out_path)
    return out_path


def get_default_modules():
    return [dict(m) for m in DEFAULT_MODULES]


def get_default_hide():
    return list(DEFAULT_HIDE)


if __name__ == '__main__':
    print('自检：')
    d, r, p = _load_analysis()
    print('数据行数:', len(d), '地区×年度:', len(r))
    generate_from_data(d, r)
    print('已生成:', REPORT_DEFAULT_DOCX)